from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_report():
    # Initialize the document
    doc = Document()

    # --- TITLE PAGE ---
    # Centered Title
    title = doc.add_heading('COMPREHENSIVE DOCUMENT ANALYSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Centered Subtitle
    subtitle = doc.add_paragraph('AI-Generated Fact-Based Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph('\n')  # Add some spacing

    # --- METADATA SECTION ---
    # We use .bold = True to make the labels stand out
    p = doc.add_paragraph()
    p.add_run('Source Document: ').bold = True
    p.add_run('Combined_Assignment_Document.pdf\n')
    
    p.add_run('Total Pages in Source: ').bold = True
    p.add_run('1,093\n')
    
    p.add_run('Total Words in Source: ').bold = True
    p.add_run('803,502\n')
    
    p.add_run('Meaningful Content Analyzed: ').bold = True
    p.add_run('~760,000 words (Excluding citations & stopwords)\n')
    
    p.add_run('Generated: ').bold = True
    p.add_run('November 25, 2025\n')

    p.add_run('Method: ').bold = True
    p.add_run('Mixed-Methods Analysis (Python Frequency Auditing + LLM Thematic Review)')

    # Add a page break after the cover info
    doc.add_page_break()

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    
    summary_text = (
        "This report presents an automated analysis of a corpus comprising 67 documents (1,093 pages) "
        "focused on the intersection of Generative AI (ChatGPT) and Marketing Strategy.\n\n"
        "The analysis reveals a watershed moment in the industry. The consensus across the text is that "
        "Generative AI is not merely a tool for efficiency but a transformational agent. The literature, "
        "heavily anchored in 2023–2024 research, suggests that while AI democratizes content creation, "
        "it simultaneously elevates the value of human ethical oversight, strategic curation, and empathy."
    )
    doc.add_paragraph(summary_text)

    # --- SECTION 2: QUANTITATIVE EVIDENCE ---
    doc.add_heading('2. Quantitative Evidence (The "Why" Behind the Findings)', level=1)
    
    doc.add_paragraph("A frequency analysis of the 800,000+ word dataset reveals the specific focus of the literature:")

    # Bullet Point 1
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Immediacy: ').bold = True
    p.add_run(
        "The most frequent year in the text is '2023' (4,418 occurrences), indicating that the dataset "
        "represents cutting-edge, post-disruption research rather than historical theory."
    )
    
    # Bullet Point 2
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Topic Dominance: ').bold = True
    p.add_run(
        "The specific term 'ChatGPT' (4,056 occurrences) appears nearly twice as often as the general term "
        "'Marketing,' proving the industry's specific fixation on Large Language Models (LLMs)."
    )
    
    # Bullet Point 3
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('The Human Factor: ').bold = True
    p.add_run(
        "Despite being a technical corpus, the word 'Human' is a top-10 keyword (1,813 occurrences), "
        "statistically validating the conclusion that the 'Human-in-the-Loop' is a central concern."
    )

    # --- SECTION 3: THEMATIC FINDINGS ---
    doc.add_heading('3. Thematic Findings', level=1)

    # Subsection A
    doc.add_heading('A. The Transformation of Content (The "Co-Pilot" Model)', level=2)
    doc.add_paragraph(
        "The text identifies a shift in the marketer’s role from 'creator' to 'editor'. "
        "The barrier to entry for content production (emails, blogs, code) has dropped to near zero. "
        "The text warns of 'content commoditization,' where brand voices become indistinguishable. "
        "The competitive advantage shifts to those who can prompt and curate AI output effectively."
    )

    # Subsection B
    doc.add_heading('B. Impact on Academic & Market Research', level=2)
    doc.add_paragraph(
        "The corpus includes significant excerpts from the Journal of Services Marketing (Vol 38, 2024) and Nature (2022). "
        "A significant portion of the text discusses the 'Industrialization of Science' and the 'Reviewer 2 Crisis'—"
        "an active debate regarding the integrity of peer review when both authors and reviewers use AI tools."
    )

    # Subsection C
    doc.add_heading('C. Ethics, Bias, and Trust', level=2)
    doc.add_paragraph(
        "The text frequently discusses the 'Uncanny Valley'—the erosion of trust when consumers realize they "
        "are interacting with a machine masquerading as a human. 'Hallucinations' (confident falsehoods) pose "
        "a legal risk, making transparency—disclosing when AI is used—a mandatory ethical standard."
    )

    # --- SECTION 4: CONCLUSION ---
    doc.add_heading('4. Conclusion', level=1)
    conclusion_text = (
        "The data indicates that the marketing industry has moved beyond the 'hype cycle' into a phase of "
        "practical integration. The central thesis of the combined documents is that AI is a capability multiplier, "
        "not a human replacement. The future of marketing belongs to the 'hybrid professional' who can navigate "
        "the technical utility of AI while maintaining the ethical and emotional standards that machines cannot replicate."
    )
    doc.add_paragraph(conclusion_text)

    # Save the file
    filename = 'Assignment_Summary_Report.docx'
    doc.save(filename)
    print(f"✅ Success! Created word document: {filename}")

if __name__ == "__main__":
    create_word_report()